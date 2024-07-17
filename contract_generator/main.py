# main.py

import streamlit as st
import pandas as pd
from zipfile import ZipFile
import os
from io import BytesIO
from docx import Document
from datetime import datetime, date
import re
import fitz  # PyMuPDF

# PDF를 이미지로 변환하는 함수
def convert_pdf_to_images(pdf_path):
    if not os.path.exists(pdf_path):
        st.error(f"PDF 파일을 찾을 수 없습니다: {pdf_path}")
        return []
    
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        img_data = pix.tobytes("png")
        images.append(img_data)
    return images

def calculate_contract_days(start_date, end_date):
    start = pd.to_datetime(start_date)
    end = pd.to_datetime(end_date)
    days = (end - start).days + 1  # 종료일도 포함
    return days

def replace_keywords(doc, keywords):
    date_fields = ['{계약시작일}', '{계약마감일}', '{납품기일}']
    number_fields = ['{지급금액}', '{납품금액}', '{상금}']
    
    for paragraph in doc.paragraphs:
        for key, value in keywords.items():
            if key in paragraph.text:
                if key in date_fields:
                    try:
                        date_value = pd.to_datetime(value).strftime('%Y-%m-%d')
                        paragraph.text = paragraph.text.replace(key, date_value)
                    except:
                        paragraph.text = paragraph.text.replace(key, str(value))
                elif key in number_fields:
                    formatted_value = format_number_with_commas(value)
                    paragraph.text = paragraph.text.replace(key, formatted_value)
                else:
                    paragraph.text = paragraph.text.replace(key, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in keywords.items():
                    if key in cell.text:
                        if key in date_fields:
                            try:
                                date_value = pd.to_datetime(value).strftime('%Y-%m-%d')
                                cell.text = cell.text.replace(key, date_value)
                            except:
                                cell.text = cell.text.replace(key, str(value))
                        elif key in number_fields:
                            formatted_value = format_number_with_commas(value)
                            cell.text = cell.text.replace(key, formatted_value)
                        else:
                            cell.text = cell.text.replace(key, str(value))

    special_keywords = {
        '{생년월일}': lambda k: convert_ssn_to_birthdate(keywords.get('{주민등록번호}', '')),
        '{오늘날짜}': lambda k: date.today().strftime('%Y-%m-%d'),
        '{납품금액한글}': lambda k: convert_number_to_korean(keywords.get('{납품금액}', '0')),
        '{상금한글}': lambda k: convert_number_to_korean(keywords.get('{상금}', '0')),
        '{일시}': lambda k: format_date_only(keywords.get('{일시}', '')),
        '{과업일자}': lambda k: format_date_only(keywords.get('{과업일자}', '')),
        '{계약시작일}': lambda k: format_date_only(keywords.get('{계약시작일}', '')),
        '{계약마감일}': lambda k: format_date_only(keywords.get('{계약마감일}', '')),
        '{근무일}': lambda k: format_work_period(
            keywords.get('{계약시작일}', ''),
            keywords.get('{계약마감일}', '')
        )
    }

    for paragraph in doc.paragraphs:
        for key, func in special_keywords.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, func(key))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, func in special_keywords.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, func(key))

def convert_ssn_to_birthdate(ssn):
    birth_date = ssn.split('-')[0]
    year = birth_date[:2]
    month = birth_date[2:4]
    day = birth_date[4:6]
    
    if int(year) < 22:  # 2000년대 출생
        year = f'20{year}'
    else:  # 1900년대 출생
        year = f'19{year}'
    
    return f'{year}.{month}.{day}'

def convert_number_to_korean(number):
    units = ["", "만", "억", "조", "경"]
    num_str = str(number)
    num_str = num_str.zfill(((len(num_str) + 3) // 4) * 4)
    result = []
    for i in range(0, len(num_str), 4):
        part = num_str[i:i+4]
        if part != "0000":
            part_korean = convert_part_to_korean(part)
            result.append(part_korean + units[(len(num_str) - i) // 4 - 1])
    return ''.join(result)

def convert_part_to_korean(part):
    digits = ["", "일", "이", "삼", "사", "오", "육", "칠", "팔", "구"]
    units = ["", "십", "백", "천"]
    result = []
    for i, digit in enumerate(part):
        if digit != "0":
            result.append(digits[int(digit)] + units[3 - i])
    return ''.join(result)

def format_number_with_commas(number):
    try:
        return f"{int(float(number)):,}"
    except ValueError:
        return str(number)

def format_date_only(datetime_str):
    return pd.to_datetime(datetime_str).strftime('%Y-%m-%d')

def format_work_period(start_date, end_date):
    try:
        start = pd.to_datetime(start_date).strftime('%Y-%m-%d')
        end = pd.to_datetime(end_date).strftime('%Y-%m-%d')
        days = calculate_contract_days(start_date, end_date)
        return f"{start} ~ {end} ({days}일간)"
    except:
        return "날짜 형식 오류"

# 파일 이름 생성을 위한 함수
def generate_filename(keywords, today):
    name = keywords.get('{이름}', 'Unknown')
    contract_type = keywords.get('{계약명}', 'Contract')
    return f"{today}_{name}_{contract_type}.docx"

def main():
    st.title("계약서 생성기")

    # 템플릿 선택
    template_options = {
        "일반 대행 용역 계약서": "General Service.docx",
        "일용직 근로자 계약서": "Temporary Worker.docx",
        "수당지급 약정서": "Allowance Payment.docx",
        "상금지급 약정서": "Bonus Payment.docx"
    }
    
    selected_template = st.selectbox(
        "사용할 템플릿을 선택하세요:",
        options=list(template_options.keys())
    )

    selected_template_file = template_options[selected_template]

    # PDF 미리보기
    pdf_path = os.path.join(os.path.dirname(__file__), f"data/{selected_template_file.split('.')[0]}.pdf")
    images = convert_pdf_to_images(pdf_path)
    if images:
        st.markdown(f"### {selected_template} 예시")
        cols = st.columns(2)
        for i, image in enumerate(images):
            with cols[i % 2]:
                st.image(image, use_column_width=True)

    # 예시 엑셀 파일 제공
    st.markdown(f"### {selected_template} 예시 엑셀 템플릿")
    example_file_path = os.path.join(os.path.dirname(__file__), f'data/{selected_template_file.split(".")[0]}_Template.xlsx')
    if not os.path.exists(example_file_path):
        st.error(f"엑셀 템플릿 파일을 찾을 수 없습니다: {example_file_path}")
    else:
        with open(example_file_path, 'rb') as file:
            st.download_button(
                label=f"{selected_template} 예시 엑셀 다운로드",
                data=file,
                file_name=f"{selected_template_file.split('.')[0]}_Template.xlsx"
            )

    # 파일 업로드
    excel_file = st.file_uploader(f"{selected_template} 엑셀 파일 업로드", type="xlsx")
    
    if excel_file:
        df = pd.read_excel(excel_file, header=0)
        df.columns = [f'{{{col}}}' for col in df.columns]

        today = date.today().strftime('%Y%m%d')
        contract_count = len(df)

        if selected_template == "일반 대행 용역 계약서":
            project_name = df['{프로젝트명}'].iloc[0]  # 첫 번째 행의 프로젝트명 사용
            folder_name = f"{project_name}_{today}_{contract_count}"
        else:
            folder_name = f"{today}_{selected_template}_{contract_count}"

        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, 'w') as zipf:
            template_path = os.path.join(os.path.dirname(__file__), f"templates/{selected_template_file}")

            for i, row in df.iterrows():
                keywords = row.to_dict()
                filled_doc = Document(template_path)

                replace_keywords(filled_doc, keywords)
                
                if selected_template == "일반 대행 용역 계약서":
                    contract_filename = f"{today}_{keywords.get('{사업자명}', 'Unknown')}_{keywords.get('{프로젝트명}', 'Unknown')}.docx"
                else:
                    contract_filename = generate_filename(keywords, today)

                doc_buffer = BytesIO()
                filled_doc.save(doc_buffer)
                doc_buffer.seek(0)
                zipf.writestr(f"{folder_name}/{contract_filename}", doc_buffer.getvalue())

        zip_buffer.seek(0)
        st.download_button('계약서 다운로드', zip_buffer, file_name=f'{folder_name}.zip')

if __name__ == "__main__":
    main()